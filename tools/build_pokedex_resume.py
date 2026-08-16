from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.colors import HexColor
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image, KeepTogether, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

OUT = Path(r"D:\AndroidApp\简历")
OUT.mkdir(parents=True, exist_ok=True)
STEM = "杨书涵-武汉工程科技学院-应届Android开发工程师"
PHOTO = Path(r"C:\Users\jy420\Pictures\微信图片_20260810112537_1_4.jpg")
FONT = "Microsoft YaHei"
BLUE, NAVY, GOLD, GRAY, PALE = "3E6D82", "203A4A", "B9965B", "5B6670", "F4F7F8"

PROJECT_TITLE = "宝可梦图鉴机（Pokédex Scanner）- 可扩展的 Android 智能图鉴应用｜独立主导开发"
PROJECT_STACK = "Kotlin、Jetpack Compose、Material 3、ViewModel、StateFlow、CameraX、协程、OkHttp、WorkManager、SharedPreferences"
PROJECT_BULLETS = [
    ("状态架构：", "构建扫描、图鉴详情、收藏、画廊、设置与小游戏等多页面流程，使用 ViewModel、StateFlow 统一界面状态与事件。"),
    ("智能识别：", "接入 CameraX 拍摄与相册导入，完成图片尺寸限制、JPEG 压缩和 Base64 编码；适配 Responses、Chat Completions、Anthropic Messages 三类接口协议及多模型预设，处理超时、取消、解析与错误分类。"),
    ("资源与语音：", "实现图鉴资源包的清单拉取、断点下载、SHA-256 校验、暂存目录切换与修复；使用 WorkManager 管理语音包后台下载、网络确认、版本检测和失败状态。"),
    ("离线与质量：", "实现本地图鉴、多媒体语音播报、音效/BGM、AI 响应历史和 5 类小游戏；项目包含 22 个测试文件，覆盖状态流转、协议解析、资源清单、语音包和游戏逻辑。"),
]

def set_font(run, size=9, bold=False, color=NAVY):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)

def hide_borders(table):
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}"); tag.set(qn("w:val"), "nil"); borders.append(tag)
    table._tbl.tblPr.append(borders)

def shade(cell, fill):
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); cell._tc.get_or_add_tcPr().append(shd)

def add_text(p, text, size=9, bold=False, color=NAVY):
    run = p.add_run(text); set_font(run, size, bold, color); return run

def section_docx(doc, title):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(3.5); p.paragraph_format.space_after = Pt(1.5); p.paragraph_format.line_spacing = 1
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), BLUE); p._p.get_or_add_pPr().append(shd)
    add_text(p, title, 9.5, True, "FFFFFF")

def bullet_docx(doc, label, text):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(.13); p.paragraph_format.first_line_indent = Inches(-.13); p.paragraph_format.space_after = Pt(.3); p.paragraph_format.line_spacing = 1.02
    add_text(p, "• ", 8.35, True, GOLD); add_text(p, label, 8.35, True); add_text(p, text, 8.35)

def project_docx(doc, title, stack, bullets, link=None):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(.2); add_text(p, title, 9.15, True)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(.2); add_text(p, "技术栈：", 8.2, True, BLUE); add_text(p, stack, 8.2, False, GRAY)
    for label, text in bullets: bullet_docx(doc, label, text)
    if link:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(.2); add_text(p, "开源地址：", 8.2, True, BLUE); add_text(p, link, 8.2, False, GRAY)

def build_docx():
    doc = Document(); sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.27), Inches(11.69)
    sec.top_margin = sec.bottom_margin = Inches(.32); sec.left_margin = sec.right_margin = Inches(.5)
    doc.styles["Normal"].font.name = FONT; doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    header = doc.add_table(rows=1, cols=2); hide_borders(header)
    left, right = header.rows[0].cells
    for cell in (left, right): shade(cell, PALE); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = left.paragraphs[0]; p.paragraph_format.space_after = Pt(0); add_text(p, "杨书涵", 21.5, True, BLUE)
    p = left.add_paragraph(); p.paragraph_format.space_after = Pt(0); add_text(p, "Android开发工程师（应届初级）", 10, True, GOLD)
    p = left.add_paragraph(); p.paragraph_format.space_after = Pt(0); add_text(p, "意向城市：无地域限制，接受全国任意城市调配  |  19186400980  |  serenesweet2022@outlook.com", 8.2, False, GRAY)
    p = left.add_paragraph(); p.paragraph_format.space_after = Pt(0); add_text(p, "GitHub：github.com/wuzixu2007  |  电子信息本科应届生，长期专注 Android 移动端开发。", 8.2, False, GRAY)
    p = right.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(str(PHOTO), width=Inches(.8), height=Inches(1.05))
    section_docx(doc, "核心技能")
    for label, value in [
        ("Android 客户端", "Kotlin、Jetpack Compose、Material 3、ViewModel、StateFlow、协程、Navigation Compose、CameraX"),
        ("网络与任务", "OkHttp、多协议 AI 接口适配、图片压缩/Base64、WorkManager 后台任务、资源下载与 SHA-256 校验"),
        ("数据与质量", "SharedPreferences、JSON 本地历史、JUnit 单元测试、Android 仪器测试、Gradle、Git/GitHub"),
        ("软硬件基础", "C 语言、ESP32、多传感器数据采集、无线通信与软硬件联调"),
    ]:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(.2); p.paragraph_format.line_spacing = 1.02; add_text(p, label + "：", 8.3, True, BLUE); add_text(p, value, 8.3)
    section_docx(doc, "重点项目经历")
    project_docx(doc, PROJECT_TITLE, PROJECT_STACK, PROJECT_BULLETS, "github.com/wuzixu2007/android-pokedex")
    section_docx(doc, "其他项目与教育")
    project_docx(doc, "AIDIIS - 编程学习类 App｜独立开发", "Kotlin、Compose、Navigation、Room、DataStore、协程", [("实现：", "完成多页面与底部导航、用户积分和学习进度本地存储、课程偏好恢复；云端接口调用结构待配置服务地址。")])
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(.2); add_text(p, "武汉工程科技学院｜电子信息工程（本科）｜2022.09 - 2026.06", 8.85, True)
    bullet_docx(doc, "ESP32 实践：", "完成温湿度、红外等传感器的数据采集、无线传输和软硬件联调。")
    section_docx(doc, "工作经历与个人优势")
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(.2); add_text(p, "北京乐博乐博教育科技有限公司｜少儿编程讲师｜2025.09 - 2026.01", 8.7, True)
    bullet_docx(doc, "", "讲授 Arduino、C++、Python 课程，承担技术逻辑拆解、程序问题排查与沟通交付。")
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(.2); add_text(p, "欣博文化培训学校｜兼职化学教师｜2023.03 - 2025.06", 8.7, True)
    bullet_docx(doc, "", "在校兼职，负责课程规划与沟通，具备稳定交付与用户沟通能力。")
    bullet_docx(doc, "", "拥有从 0 到 1 的 Android 项目实践，重视状态管理、异常处理、可维护性与用户体验。")
    footer = sec.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER; add_text(footer, "杨书涵｜应届 Android 开发工程师｜19186400980", 7.1, False, GRAY)
    doc.save(OUT / f"{STEM}.docx")

def build_pdf():
    pdfmetrics.registerFont(TTFont("CN", r"C:\Windows\Fonts\Deng.ttf")); pdfmetrics.registerFont(TTFont("CNB", r"C:\Windows\Fonts\Dengb.ttf"))
    styles = {
        "title": ParagraphStyle("title", fontName="CNB", fontSize=21, leading=23, textColor=HexColor("#"+BLUE)),
        "job": ParagraphStyle("job", fontName="CNB", fontSize=9.7, leading=11, textColor=HexColor("#"+GOLD)),
        "body": ParagraphStyle("body", fontName="CN", fontSize=8.8, leading=11.6, textColor=HexColor("#"+NAVY), spaceAfter=.7),
        "meta": ParagraphStyle("meta", fontName="CN", fontSize=8.15, leading=10.5, textColor=HexColor("#"+GRAY)),
        "project": ParagraphStyle("project", fontName="CNB", fontSize=9.55, leading=12.2, textColor=HexColor("#"+NAVY), spaceAfter=.5),
    }
    def p(text, style="body"): return Paragraph(text, styles[style])
    def section(title):
        t = Table([[p(title)]], colWidths=[182*mm]); t.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,-1),HexColor("#"+BLUE)),("TEXTCOLOR",(0,0),(-1,-1),colors.white),("LEFTPADDING",(0,0),(-1,-1),5),("TOPPADDING",(0,0),(-1,-1),3),("BOTTOMPADDING",(0,0),(-1,-1),3),("LINEBELOW",(0,0),(-1,-1),.8,HexColor("#"+GOLD))])); return [Spacer(1,4),t,Spacer(1,2.5)]
    def bullet(label, text): return p(f'<font color="#{GOLD}"><b>• </b></font><b>{label}</b>{text}')
    def project(title, stack, bullets, link=None):
        items=[p(title,"project"),p(f'<font color="#{BLUE}"><b>技术栈：</b></font><font color="#{GRAY}">{stack}</font>',"meta")]+[bullet(a,b) for a,b in bullets]
        if link: items.append(p(f'<font color="#{BLUE}"><b>开源地址：</b></font><font color="#{GRAY}">{link}</font>',"meta"))
        return KeepTogether(items)
    flow=[]
    profile=[p("杨书涵","title"),p("Android开发工程师（应届初级）","job"),p("意向城市：无地域限制，接受全国任意城市调配  |  19186400980  |  serenesweet2022@outlook.com","meta"),p("GitHub：github.com/wuzixu2007  |  电子信息本科应届生，长期专注 Android 移动端开发。","meta")]
    head=Table([[profile,Image(str(PHOTO),width=21*mm,height=28*mm)]],colWidths=[152*mm,30*mm],rowHeights=[30*mm]); head.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,-1),HexColor("#"+PALE)),("VALIGN",(0,0),(-1,-1),"MIDDLE"),("LEFTPADDING",(0,0),(-1,-1),5),("RIGHTPADDING",(0,0),(-1,-1),5),("TOPPADDING",(0,0),(-1,-1),2),("BOTTOMPADDING",(0,0),(-1,-1),2),("ALIGN",(1,0),(1,0),"CENTER")]))
    flow.append(head); flow+=section("核心技能")
    for a,b in [("Android 客户端","Kotlin、Jetpack Compose、Material 3、ViewModel、StateFlow、协程、Navigation Compose、CameraX"),("网络与任务","OkHttp、多协议 AI 接口适配、图片压缩/Base64、WorkManager 后台任务、资源下载与 SHA-256 校验"),("数据与质量","SharedPreferences、JSON 本地历史、JUnit 单元测试、Android 仪器测试、Gradle、Git/GitHub"),("软硬件基础","C 语言、ESP32、多传感器数据采集、无线通信与软硬件联调")]: flow.append(p(f'<font color="#{BLUE}"><b>{a}：</b></font>{b}'))
    flow+=section("重点项目经历"); flow.append(project(PROJECT_TITLE,PROJECT_STACK,PROJECT_BULLETS,"github.com/wuzixu2007/android-pokedex"))
    flow+=section("其他项目与教育")
    flow.append(project("AIDIIS - 编程学习类 App｜独立开发","Kotlin、Compose、Navigation、Room、DataStore、协程",[("实现：","完成多页面与底部导航、用户积分和学习进度本地存储、课程偏好恢复；云端接口调用结构待配置服务地址。")]))
    flow += [p("<b>武汉工程科技学院｜电子信息工程（本科）｜2022.09 - 2026.06</b>"),bullet("ESP32 实践：","完成温湿度、红外等传感器的数据采集、无线传输和软硬件联调。")]
    flow+=section("工作经历与个人优势")
    flow += [p("<b>北京乐博乐博教育科技有限公司｜少儿编程讲师｜2025.09 - 2026.01</b>"),bullet("","讲授 Arduino、C++、Python 课程，承担技术逻辑拆解、程序问题排查与沟通交付。"),p("<b>欣博文化培训学校｜兼职化学教师｜2023.03 - 2025.06</b>"),bullet("","在校兼职，负责课程规划与沟通，具备稳定交付与用户沟通能力。"),bullet("","拥有从 0 到 1 的 Android 项目实践，重视状态管理、异常处理、可维护性与用户体验。")]
    def footer(canvas, _):
        canvas.saveState(); canvas.setFont("CN",7); canvas.setFillColor(HexColor("#"+GRAY)); canvas.drawCentredString(A4[0]/2,6*mm,"杨书涵｜应届 Android 开发工程师｜19186400980"); canvas.restoreState()
    SimpleDocTemplate(str(OUT/f"{STEM}.pdf"),pagesize=A4,leftMargin=14*mm,rightMargin=14*mm,topMargin=9*mm,bottomMargin=10*mm).build(flow,onFirstPage=footer,onLaterPages=footer)

if __name__ == "__main__":
    build_docx(); build_pdf()
